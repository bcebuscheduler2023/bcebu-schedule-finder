import streamlit as st
import pandas as pd
import docx
import pdfplumber
import json
import os
import re

st.set_page_config(page_title="B'Cebu Class Schedule Finder", page_icon="📚", layout="wide")

st.title("🎓 B'Cebu Class Schedule System")
st.caption("API NEXT EDU, Inc. — Class Schedule Search & Extractor App")

DB_FILE = "schedules_data.json"

def load_schedules():
    if os.path.exists(DB_FILE):
        try:
            with open(DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []

def save_schedules(data):
    with open(DB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

if "schedules_db" not in st.session_state:
    st.session_state["schedules_db"] = load_schedules()

def parse_docx(file):
    doc = docx.Document(file)
    raw_texts = []
    
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            raw_texts.append(t)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip().replace('\n', ' ')
                if txt and (not raw_texts or raw_texts[-1] != txt):
                    raw_texts.append(txt)

    return process_extracted_texts(raw_texts, doc.tables, file.name)

def parse_pdf(file):
    raw_texts = []
    extracted_rows = []
    
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    clean_row = [str(c).strip().replace('\n', ' ') for c in row if c is not None and str(c).strip()]
                    if clean_row:
                        extracted_rows.append(clean_row)
                        for cell_text in clean_row:
                            if not raw_texts or raw_texts[-1] != cell_text:
                                raw_texts.append(cell_text)
                                
    return process_pdf_data(raw_texts, extracted_rows, file.name)

def process_extracted_texts(raw_texts, tables, filename):
    def find_field(keyword):
        for idx, item in enumerate(raw_texts):
            clean_item = item.lower().replace(":", "").strip()
            if clean_item == keyword.lower():
                if idx + 1 < len(raw_texts):
                    val = raw_texts[idx + 1].replace("|", "").strip()
                    if val and not any(k in val.lower() for k in ["name", "nickname", "course", "duration", "nationality"]):
                        return val
        return "N/A"

    name = find_field("Name")
    nickname = find_field("Nickname")
    course = find_field("Course")
    duration = find_field("Duration")
    nationality = find_field("Nationality")

    clean_filename = os.path.splitext(filename)[0]
    final_name = name if name != "N/A" else f"Student ({clean_filename})"

    student_data = {
        "Name": final_name,
        "Nickname": nickname if nickname != "N/A" else "-",
        "Course": course if course != "N/A" else "-",
        "Duration": duration if duration != "N/A" else "-",
        "Nationality": nationality if nationality != "N/A" else "-",
        "Schedule": []
    }

    for table in tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                txt = cell.text.strip().replace('\n', ' ')
                if txt and (not row_cells or row_cells[-1] != txt):
                    row_cells.append(txt)
            
            if row_cells and re.search(r"\d{2}:\d{2}\s*-\s*\d{2}:\d{2}", row_cells[0]):
                student_data["Schedule"].append({
                    "Time Period": row_cells[0] if len(row_cells) > 0 else "-",
                    "Period": row_cells[1] if len(row_cells) > 1 else "-",
                    "Room": row_cells[2] if len(row_cells) > 2 else "-",
                    "Teacher": row_cells[3] if len(row_cells) > 3 else "-",
                    "Type": row_cells[4] if len(row_cells) > 4 else "-",
                    "Subject / Remarks": row_cells[5] if len(row_cells) > 5 else (row_cells[-1] if len(row_cells) > 3 else "-")
                })

    return student_data

def process_pdf_data(raw_texts, rows, filename):
    def find_field(keyword):
        for idx, item in enumerate(raw_texts):
            clean_item = item.lower().replace(":", "").strip()
            if clean_item == keyword.lower():
                if idx + 1 < len(raw_texts):
                    val = raw_texts[idx + 1].replace("|", "").strip()
                    if val and not any(k in val.lower() for k in ["name", "nickname", "course", "duration", "nationality"]):
                        return val
        return "N/A"

    name = find_field("Name")
    nickname = find_field("Nickname")
    course = find_field("Course")
    duration = find_field("Duration")
    nationality = find_field("Nationality")

    clean_filename = os.path.splitext(filename)[0]
    final_name = name if name != "N/A" else f"Student ({clean_filename})"

    student_data = {
        "Name": final_name,
        "Nickname": nickname if nickname != "N/A" else "-",
        "Course": course if course != "N/A" else "-",
        "Duration": duration if duration != "N/A" else "-",
        "Nationality": nationality if nationality != "N/A" else "-",
        "Schedule": []
    }

    for row_cells in rows:
        if row_cells and re.search(r"\d{2}:\d{2}\s*-\s*\d{2}:\d{2}", row_cells[0]):
            student_data["Schedule"].append({
                "Time Period": row_cells[0] if len(row_cells) > 0 else "-",
                "Period": row_cells[1] if len(row_cells) > 1 else "-",
                "Room": row_cells[2] if len(row_cells) > 2 else "-",
                "Teacher": row_cells[3] if len(row_cells) > 3 else "-",
                "Type": row_cells[4] if len(row_cells) > 4 else "-",
                "Subject / Remarks": row_cells[5] if len(row_cells) > 5 else (row_cells[-1] if len(row_cells) > 3 else "-")
            })

    return student_data

# App Navigation
tab_search, tab_upload, tab_records = st.tabs(["🔍 Search Student Schedules", "📤 Admin Upload Files (.docx / .pdf)", "📋 Stored Records"])

# TAB 1: SEARCH INTERFACE
with tab_search:
    st.subheader("Search Schedules")
    total_students = len(st.session_state["schedules_db"])
    st.caption(f"📊 Currently storing **{total_students}** student schedule(s).")
    
    search_query = st.text_input("Enter Student Full Name or Nickname:", "").strip()

    if search_query:
        matches = [
            s for s in st.session_state["schedules_db"]
            if search_query.lower() in s["Name"].lower() or search_query.lower() in s["Nickname"].lower()
        ]

        if matches:
            for student in matches:
                with st.expander(f"📌 {student['Name']} ({student['Nickname']}) — {student['Course']}", expanded=True):
                    c1, c2 = st.columns(2)
                    c1.write(f"**Nationality:** {student['Nationality']}")
                    c2.write(f"**Duration:** {student['Duration']}")
                    st.divider()
                    
                    st.markdown("#### ⏰ Class Schedule")
                    if student["Schedule"]:
                        st.table(pd.DataFrame(student["Schedule"]))
                    else:
                        st.warning("No individual class periods detected in this document.")
        else:
            st.warning(f"No student found matching '{search_query}'.")
    else:
        st.info("💡 Type a student name or nickname above to view their schedule.")

# TAB 2: UPLOAD INTERFACE
with tab_upload:
    st.subheader("Upload Word Documents (.docx) or PDFs (.pdf)")
    st.caption("Select one or multiple student schedule files to import them into the search system.")
    
    uploaded_files = st.file_uploader("Drop .docx or .pdf files here", type=["docx", "pdf"], accept_multiple_files=True)

    if uploaded_files:
        if st.button("Extract & Save Schedules"):
            count = 0
            for file in uploaded_files:
                try:
                    if file.name.endswith(".pdf"):
                        parsed_info = parse_pdf(file)
                    else:
                        parsed_info = parse_docx(file)
                    
                    st.session_state["schedules_db"] = [
                        s for s in st.session_state["schedules_db"] 
                        if s["Name"].lower() != parsed_info["Name"].lower()
                    ]
                    
                    st.session_state["schedules_db"].append(parsed_info)
                    count += 1
                except Exception as e:
                    st.error(f"Error reading {file.name}: {e}")

            save_schedules(st.session_state["schedules_db"])
            st.success(f"Successfully processed {count} file(s)! Check the 'Stored Records' or 'Search' tab.")

# TAB 3: RECORDS LIST
with tab_records:
    st.subheader("List of All Stored Students")
    if st.session_state["schedules_db"]:
        student_list = [
            {
                "Full Name / Record ID": s["Name"], 
                "Nickname": s["Nickname"], 
                "Course": s["Course"],
                "Classes Extracted": len(s["Schedule"])
            } 
            for s in st.session_state["schedules_db"]
        ]
        st.table(pd.DataFrame(student_list))
        
        if st.button("Clear All Stored Schedules"):
            st.session_state["schedules_db"] = []
            save_schedules([])
            st.rerun()
    else:
        st.info("No records saved in database.")
